import io
import os
import re
from datetime import datetime
from django.conf import settings
from django.template.loader import render_to_string
from django.utils.timezone import now
from django.utils.safestring import mark_safe

# PDF
from xhtml2pdf import pisa

# DOCX
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# XLSX
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

from testcases.models import TestCase

def format_text_pdf(text):
    if not text:
        return ""
    text = str(text)
    # Normalize unicode hyphens/dashes to standard ASCII hyphen-minus
    unicode_dashes = ['\u2010', '\u2011', '\u2012', '\u2013', '\u2014', '\u2015', '\u2212']
    for dash in unicode_dashes:
        text = text.replace(dash, '-')
    # Add a space after a hyphen ONLY if followed by a letter or number
    return re.sub(r'-(?=\w)', '- ', text)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def style_table_header(row, col_widths, alignments=None, header_shading="1F4E78"):
    for idx, cell in enumerate(row.cells):
        if col_widths and idx < len(col_widths):
            cell.width = col_widths[idx]
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        set_cell_shading(cell, header_shading)
        p = cell.paragraphs[0]
        if alignments and idx < len(alignments):
            if alignments[idx] == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignments[idx] == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)


def add_styled_row(table, data, col_widths=None, alignments=None):
    row_cells = table.add_row().cells
    for idx, val in enumerate(data):
        cell = row_cells[idx]
        if col_widths and idx < len(col_widths):
            cell.width = col_widths[idx]
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        p.text = str(val) if val is not None else ""
        
        if alignments and idx < len(alignments):
            if alignments[idx] == 'center':
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignments[idx] == 'right':
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(9.5)


def add_styled_heading(document, text, level):
    h = document.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    for run in h.runs:
        run.font.name = 'Arial'
        if level == 1:
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(31, 78, 120)  # #1F4E78
        else:
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(52, 73, 94)
    return h


class ReportEngine:
    @staticmethod
    def get_company_info():
        return {
            "name": "Observatory Management System",
            "tagline": "Enterprise SDLC Governance & Traceability",
            "logo_path": os.path.join(settings.BASE_DIR, "static/img/company_logo.png"),
            "confidential_label": "",  # Public distribution
        }

    # ==========================================
    # MASTER REPORT GENERATOR
    # ==========================================
    @staticmethod
    def generate_master_report(project, format='pdf'):
        """Generates a comprehensive master report including Summary, Team, Req, Tasks, TC, and RTM"""
        requirements = project.requirements.filter(is_in_trash=False)
        tasks = project.tasks.filter(is_in_trash=False)
        test_cases = project.test_cases.filter(is_in_trash=False)
        
        if format in ['pdf', 'html']:
            return ReportEngine._generate_master_pdf_html(project, requirements, tasks, test_cases, format)
        elif format == 'docx':
            return ReportEngine._generate_master_docx(project, requirements, tasks, test_cases)
        elif format == 'xlsx':
            return ReportEngine._generate_master_xlsx(project, requirements, tasks, test_cases)
        elif format == 'md':
            return ReportEngine._generate_master_markdown(project, requirements, tasks, test_cases)

    @staticmethod
    def _get_rtm_data(requirements):
        rtm_data = []
        for req in requirements:
            related_tasks = req.tasks.filter(is_in_trash=False)
            related_tcs = TestCase.objects.filter(task__in=related_tasks, is_in_trash=False)
            rtm_data.append({"req": req, "tasks": related_tasks, "test_cases": related_tcs})
        return rtm_data

    @staticmethod
    def _generate_master_pdf_html(project, requirements, tasks, test_cases, format):
        # Precompute values to avoid complex template rendering
        for req in requirements:
            req.req_id_pdf = format_text_pdf(req.req_id)
            req.name_pdf = format_text_pdf(req.name)
            
            linked_tasks_list = [t.task_id for t in req.tasks.filter(is_in_trash=False)]
            req.linked_tasks_str = format_text_pdf(", ".join(linked_tasks_list) or "None")
            
            tcs = sorted(list(set([
                tc.test_id 
                for t in req.tasks.filter(is_in_trash=False) 
                for tc in t.test_cases.filter(is_in_trash=False)
            ])))
            req.linked_tcs_str = format_text_pdf(", ".join(tcs) or "None")

        for task in tasks:
            task.task_id_pdf = format_text_pdf(task.task_id)
            task.title_pdf = format_text_pdf(task.title)
            task.linked_req_str = format_text_pdf(task.requirement.req_id if task.requirement else "None")
            task.linked_tcs_str = format_text_pdf(", ".join([tc.test_id for tc in task.test_cases.filter(is_in_trash=False)]) or "None")
            task.assignees_str = ", ".join([a.display_name for a in task.assignees.all()]) or "Unassigned"
            task.due_date_str = task.due_date.strftime('%Y-%m-%d') if task.due_date else "N/A"

        for tc in test_cases:
            tc.test_id_pdf = format_text_pdf(tc.test_id)
            tc.title_pdf = format_text_pdf(tc.title)
            tc.linked_task_str = format_text_pdf(tc.task.task_id if tc.task else "None")
            tc.linked_req_str = format_text_pdf(tc.task.requirement.req_id if (tc.task and tc.task.requirement) else "None")
            tc.expected_result_pdf = format_text_pdf(tc.expected_result or "None")

        rtm_data_pdf = []
        for row in ReportEngine._get_rtm_data(requirements):
            formatted_tasks = []
            for t in row['tasks']:
                formatted_tasks.append({
                    "task_id": format_text_pdf(t.task_id),
                    "title": format_text_pdf(t.title),
                    "status_display": t.get_status_display()
                })
            
            formatted_tcs = []
            for tc in row['test_cases']:
                formatted_tcs.append({
                    "test_id": format_text_pdf(tc.test_id),
                    "status_display": tc.get_status_display()
                })
            
            rtm_data_pdf.append({
                "req": {
                    "req_id": format_text_pdf(row['req'].req_id),
                    "name": format_text_pdf(row['req'].name)
                },
                "tasks": formatted_tasks,
                "test_cases": formatted_tcs
            })

        context = {
            "company": ReportEngine.get_company_info(),
            "project": project,
            "title": f"Comprehensive Master Project Dossier - {project.name}",
            "generated_by": "System Documentation Engine",
            "timestamp": now(),
            "requirements": requirements,
            "tasks": tasks,
            "test_cases": test_cases,
            "rtm_data_pdf": rtm_data_pdf,
            "include_requirements": True,
            "include_tasks": True,
            "include_test_cases": True,
            "include_rtm": True,
        }
        html_content = render_to_string("reports/master_report_template.html", context)
        
        if format == 'html':
            return html_content.encode('utf-8')
            
        # PDF Generation
        result = io.BytesIO()
        pisa_status = pisa.CreatePDF(src=html_content, dest=result)
        if pisa_status.err:
            raise Exception("PDF generation failed via xhtml2pdf.")
        return result.getvalue()

    @staticmethod
    def _generate_master_docx(project, requirements, tasks, test_cases):
        document = Document()
        company = ReportEngine.get_company_info()
        title = f"Comprehensive Master Project Dossier - {project.name}"
        
        # Margins & Section Setup
        section = document.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Header/Footer
        header = section.header
        header.paragraphs[0].text = f"{company['name']} | {project.project_id}"
        header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer = section.footer
        footer.paragraphs[0].text = f"Report: {title} | Generated: {now().strftime('%Y-%m-%d %H:%M')}"
        footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title Page
        document.add_heading(company['name'], 0)
        document.add_heading(title, 1)
        document.add_paragraph(f"Project ID: {project.project_id}")
        document.add_paragraph(f"Status: {project.get_status_display()}")
        document.add_paragraph(f"Date: {now().strftime('%Y-%m-%d')}")
        if company.get('confidential_label'):
            document.add_paragraph(company['confidential_label'])
        document.add_page_break()

        # 1. Project Summary
        add_styled_heading(document, "1. Project Summary", level=1)
        p = document.add_paragraph(project.description or "No description provided.")
        p.paragraph_format.space_after = Pt(12)

        # 2. Team
        add_styled_heading(document, "2. Project Team", level=1)
        team_table = document.add_table(rows=1, cols=3)
        team_table.style = 'Table Grid'
        col_widths = [Inches(2.0), Inches(2.0), Inches(3.0)]
        alignments = ['left', 'left', 'left']
        hdr_cells = team_table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Role'
        hdr_cells[2].text = 'Email'
        style_table_header(team_table.rows[0], col_widths, alignments)
        
        for m in project.managers.all():
            add_styled_row(team_table, [m.display_name, 'Project Manager', m.email], col_widths, alignments)
        for m in project.members.all():
            add_styled_row(team_table, [m.display_name, 'Team Member', m.email], col_widths, alignments)

        # 3. Requirements
        add_styled_heading(document, "3. Requirements Specification", level=1)
        req_table = document.add_table(rows=1, cols=7)
        req_table.style = 'Table Grid'
        col_widths_req = [Inches(1.0), Inches(1.8), Inches(0.9), Inches(0.7), Inches(0.8), Inches(0.9), Inches(0.9)]
        aligns_req = ['center', 'left', 'center', 'center', 'center', 'left', 'left']
        hdr_cells = req_table.rows[0].cells
        hdr_cells[0].text = 'REQ ID'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Type'
        hdr_cells[3].text = 'Priority'
        hdr_cells[4].text = 'Status'
        hdr_cells[5].text = 'Linked Tasks'
        hdr_cells[6].text = 'Linked Test Cases'
        style_table_header(req_table.rows[0], col_widths_req, aligns_req)
        
        for req in requirements:
            linked_tasks_list = [t.task_id for t in req.tasks.filter(is_in_trash=False)]
            linked_tasks_str = ", ".join(linked_tasks_list) or "None"
            
            linked_tcs_list = sorted(list(set([
                tc.test_id 
                for t in req.tasks.filter(is_in_trash=False) 
                for tc in t.test_cases.filter(is_in_trash=False)
            ])))
            linked_tcs_str = ", ".join(linked_tcs_list) or "None"

            add_styled_row(req_table, [
                req.req_id, 
                req.name, 
                req.get_requirement_type_display(), 
                req.get_priority_display(),
                req.get_status_display(),
                linked_tasks_str,
                linked_tcs_str
            ], col_widths_req, aligns_req)

        # 4. Tasks
        add_styled_heading(document, "4. Development Tasks", level=1)
        task_table = document.add_table(rows=1, cols=8)
        task_table.style = 'Table Grid'
        col_widths_task = [Inches(0.9), Inches(1.5), Inches(0.9), Inches(1.0), Inches(0.7), Inches(0.6), Inches(0.8), Inches(0.6)]
        aligns_task = ['center', 'left', 'center', 'left', 'center', 'center', 'left', 'center']
        hdr_cells = task_table.rows[0].cells
        hdr_cells[0].text = 'TASK ID'
        hdr_cells[1].text = 'Title'
        hdr_cells[2].text = 'Linked REQ'
        hdr_cells[3].text = 'Assignees'
        hdr_cells[4].text = 'Status'
        hdr_cells[5].text = 'Priority'
        hdr_cells[6].text = 'Linked Test Cases'
        hdr_cells[7].text = 'Due Date'
        style_table_header(task_table.rows[0], col_widths_task, aligns_task)
        
        for task in tasks:
            assignees = ", ".join([a.display_name for a in task.assignees.all()]) or "Unassigned"
            linked_req = task.requirement.req_id if task.requirement else "None"
            linked_tcs = ", ".join([tc.test_id for tc in task.test_cases.filter(is_in_trash=False)]) or "None"
            due_date_str = task.due_date.strftime('%Y-%m-%d') if task.due_date else "N/A"
            
            add_styled_row(task_table, [
                task.task_id, 
                task.title, 
                linked_req,
                assignees, 
                task.get_status_display(), 
                task.get_priority_display(),
                linked_tcs,
                due_date_str
            ], col_widths_task, aligns_task)

        # 5. QA & Test Cases
        add_styled_heading(document, "5. QA & Test Validation", level=1)
        tc_table = document.add_table(rows=1, cols=7)
        tc_table.style = 'Table Grid'
        col_widths_tc = [Inches(1.0), Inches(1.6), Inches(0.9), Inches(0.9), Inches(0.6), Inches(0.8), Inches(1.2)]
        aligns_tc = ['center', 'left', 'center', 'center', 'center', 'center', 'left']
        hdr_cells = tc_table.rows[0].cells
        hdr_cells[0].text = 'TEST ID'
        hdr_cells[1].text = 'Scenario/Title'
        hdr_cells[2].text = 'Linked Task'
        hdr_cells[3].text = 'Linked REQ'
        hdr_cells[4].text = 'Priority'
        hdr_cells[5].text = 'Status'
        hdr_cells[6].text = 'Expected Result'
        style_table_header(tc_table.rows[0], col_widths_tc, aligns_tc)
        
        for tc in test_cases:
            linked_task = tc.task.task_id if tc.task else "None"
            linked_req = tc.task.requirement.req_id if (tc.task and tc.task.requirement) else "None"
            expected_res = tc.expected_result or "None"
            
            add_styled_row(tc_table, [
                tc.test_id, 
                tc.title, 
                linked_task,
                linked_req,
                tc.priority.upper() if tc.priority else 'N/A', 
                tc.get_status_display(),
                expected_res
            ], col_widths_tc, aligns_tc)

        document.add_page_break()

        # 6. RTM
        add_styled_heading(document, "6. Traceability Matrix (RTM)", level=1)
        rtm_data = ReportEngine._get_rtm_data(requirements)
        rtm_table = document.add_table(rows=1, cols=3)
        rtm_table.style = 'Table Grid'
        col_widths_rtm = [Inches(2.0), Inches(2.5), Inches(2.5)]
        aligns_rtm = ['left', 'left', 'left']
        hdr_cells = rtm_table.rows[0].cells
        hdr_cells[0].text = 'Requirement'
        hdr_cells[1].text = 'Tasks'
        hdr_cells[2].text = 'Test Cases'
        style_table_header(rtm_table.rows[0], col_widths_rtm, aligns_rtm)
        
        for row_data in rtm_data:
            req_text = f"{row_data['req'].req_id}\n{row_data['req'].name}"
            tasks_text = "\n".join([f"{t.task_id} ({t.get_status_display()})" for t in row_data['tasks']]) or "None"
            tcs_text = "\n".join([f"{tc.test_id} ({tc.get_status_display()})" for tc in row_data['test_cases']]) or "None"
            add_styled_row(rtm_table, [req_text, tasks_text, tcs_text], col_widths_rtm, aligns_rtm)

        f = io.BytesIO()
        document.save(f)
        return f.getvalue()

    @staticmethod
    def _generate_master_xlsx(project, requirements, tasks, test_cases):
        wb = Workbook()
        
        # Summary Sheet
        ws_sum = wb.active
        ws_sum.title = "Project Summary"
        ws_sum.append(["Project ID", project.project_id])
        ws_sum.append(["Name", project.name])
        ws_sum.append(["Status", project.get_status_display()])
        ws_sum.append(["Start Date", project.start_date.strftime('%Y-%m-%d') if project.start_date else ''])
        ws_sum.append(["End Date", project.end_date.strftime('%Y-%m-%d') if project.end_date else ''])
        ws_sum.append(["Progress", f"{project.progress}%"])

        # Req Sheet
        ws_req = wb.create_sheet(title="Requirements")
        ws_req.append(["REQ ID", "Name", "Type", "Priority", "Status", "Linked Tasks", "Linked Test Cases", "Description"])
        for req in requirements:
            linked_tasks_list = [t.task_id for t in req.tasks.filter(is_in_trash=False)]
            linked_tasks_str = ", ".join(linked_tasks_list) or "None"
            
            linked_tcs_list = sorted(list(set([
                tc.test_id 
                for t in req.tasks.filter(is_in_trash=False) 
                for tc in t.test_cases.filter(is_in_trash=False)
            ])))
            linked_tcs_str = ", ".join(linked_tcs_list) or "None"
            
            ws_req.append([
                req.req_id, 
                req.name, 
                req.get_requirement_type_display(), 
                req.get_priority_display(), 
                req.get_status_display(), 
                linked_tasks_str,
                linked_tcs_str,
                req.description
            ])
        
        # Tasks Sheet
        ws_task = wb.create_sheet(title="Tasks")
        ws_task.append(["TASK ID", "Title", "Linked REQ", "Assignees", "Status", "Priority", "Linked Test Cases", "Due Date"])
        for task in tasks:
            linked_req = task.requirement.req_id if task.requirement else "None"
            linked_tcs = ", ".join([tc.test_id for tc in task.test_cases.filter(is_in_trash=False)]) or "None"
            ws_task.append([
                task.task_id, 
                task.title, 
                linked_req,
                ", ".join([a.display_name for a in task.assignees.all()]),
                task.get_status_display(), 
                task.get_priority_display(), 
                linked_tcs,
                task.due_date.strftime('%Y-%m-%d') if task.due_date else ""
            ])
            
        # Test Cases Sheet
        ws_tc = wb.create_sheet(title="Test Cases")
        ws_tc.append(["TEST ID", "Title", "Linked TASK", "Linked REQ", "Status", "Priority", "Expected Result"])
        for tc in test_cases:
            linked_task = tc.task.task_id if tc.task else "None"
            linked_req = tc.task.requirement.req_id if (tc.task and tc.task.requirement) else "None"
            ws_tc.append([
                tc.test_id, 
                tc.title, 
                linked_task, 
                linked_req,
                tc.get_status_display(), 
                tc.priority, 
                tc.expected_result
            ])

        # RTM Sheet
        ws_rtm = wb.create_sheet(title="RTM")
        ws_rtm.append(["Requirement ID", "Requirement Name", "Task ID", "Task Status", "Test Case ID", "Test Case Status"])
        rtm_data = ReportEngine._get_rtm_data(requirements)
        for r in rtm_data:
            req = r['req']
            if not r['tasks']:
                ws_rtm.append([req.req_id, req.name, "", "", "", ""])
            else:
                for t in r['tasks']:
                    tcs = [tc for tc in r['test_cases'] if tc.task == t]
                    if not tcs:
                        ws_rtm.append([req.req_id, req.name, t.task_id, t.get_status_display(), "", ""])
                    else:
                        for tc in tcs:
                            ws_rtm.append([req.req_id, req.name, t.task_id, t.get_status_display(), tc.test_id, tc.get_status_display()])

        # Professional borders and alignments
        thin_border = Border(
            left=Side(style='thin', color='D9D9D9'),
            right=Side(style='thin', color='D9D9D9'),
            top=Side(style='thin', color='D9D9D9'),
            bottom=Side(style='thin', color='D9D9D9')
        )
        
        # Style Project Summary specifically
        ws_sum.column_dimensions['A'].width = 25
        ws_sum.column_dimensions['B'].width = 60
        for row in ws_sum.iter_rows(min_row=1, max_row=ws_sum.max_row, min_col=1, max_col=2):
            for cell in row:
                cell.font = Font(name='Segoe UI', size=10)
                cell.border = thin_border
                if cell.column == 1:
                    cell.font = Font(name='Segoe UI', size=10, bold=True)
                    cell.fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                    cell.alignment = Alignment(horizontal="left", vertical="center")
                else:
                    cell.alignment = Alignment(horizontal="left", vertical="center")

        # Style other worksheets (tabular)
        for ws in wb.worksheets:
            if ws.title == "Project Summary":
                continue
                
            ws.views.sheetView[0].showGridLines = True
            
            # Style header row (Row 1)
            for cell in ws[1]:
                cell.font = Font(name='Segoe UI', size=11, bold=True, color="FFFFFF")
                cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = thin_border
                
            # Style data rows
            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=ws.max_column):
                for cell in row:
                    cell.font = Font(name='Segoe UI', size=10)
                    cell.border = thin_border
                    
                    header_val = str(ws.cell(row=1, column=cell.column).value).lower()
                    if any(k in header_val for k in ['id', 'status', 'priority', 'date', 'progress', 'type', 'est. hr']):
                        cell.alignment = Alignment(horizontal="center", vertical="top")
                    else:
                        cell.alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
            
            # Auto-adjust column width
            for col in ws.columns:
                max_length = 0
                column_letter = col[0].column_letter
                for cell in col:
                    try:
                        if cell.value:
                            lines = str(cell.value).split('\n')
                            for line in lines:
                                if len(line) > max_length:
                                    max_length = len(line)
                    except:
                        pass
                adjusted_width = min(max(max_length + 3, 12), 60)
                ws.column_dimensions[column_letter].width = adjusted_width

        f = io.BytesIO()
        wb.save(f)
        return f.getvalue()

    @staticmethod
    def _generate_master_markdown(project, requirements, tasks, test_cases):
        lines = []
        lines.append(f"# Project Master Dossier: {project.name}")
        lines.append(f"**Project ID:** {project.project_id}  ")
        lines.append(f"**Status:** {project.get_status_display()}  ")
        lines.append(f"**Progress:** {project.progress}%  ")
        lines.append(f"**Start Date:** {project.start_date.strftime('%Y-%m-%d') if project.start_date else 'N/A'}  ")
        lines.append(f"**End Date:** {project.end_date.strftime('%Y-%m-%d') if project.end_date else 'N/A'}  ")
        lines.append("")
        lines.append("## 1. Project Summary")
        lines.append(project.description or "No description provided.")
        lines.append("")
        
        # 2. Team
        lines.append("## 2. Project Team")
        lines.append("| Name | Role | Email |")
        lines.append("| :--- | :--- | :--- |")
        for m in project.managers.all():
            lines.append(f"| {m.display_name} | Project Manager | {m.email} |")
        for m in project.members.all():
            lines.append(f"| {m.display_name} | Team Member | {m.email} |")
        lines.append("")
        
        # 3. Requirements
        if requirements.exists():
            lines.append("## 3. Requirements Specification")
            lines.append("| REQ ID | Title | Type | Priority | Status | Linked Tasks | Linked Test Cases |")
            lines.append("| :---: | :--- | :--- | :---: | :---: | :--- | :--- |")
            for req in requirements:
                linked_tasks_list = [t.task_id for t in req.tasks.filter(is_in_trash=False)]
                linked_tasks_str = ", ".join(linked_tasks_list) or "None"
                
                linked_tcs_list = sorted(list(set([
                    tc.test_id 
                    for t in req.tasks.filter(is_in_trash=False) 
                    for tc in t.test_cases.filter(is_in_trash=False)
                ])))
                linked_tcs_str = ", ".join(linked_tcs_list) or "None"
                lines.append(f"| {req.req_id} | {req.name} | {req.get_requirement_type_display()} | {req.get_priority_display()} | {req.get_status_display()} | {linked_tasks_str} | {linked_tcs_str} |")
            lines.append("")
            
        # 4. Tasks
        if tasks.exists():
            lines.append("## 4. Development Tasks")
            lines.append("| TASK ID | Title | Linked REQ | Assignees | Status | Priority | Linked Test Cases | Due Date |")
            lines.append("| :---: | :--- | :---: | :--- | :---: | :---: | :--- | :---: |")
            for task in tasks:
                assignees = ", ".join([a.display_name for a in task.assignees.all()]) or "Unassigned"
                linked_req = task.requirement.req_id if task.requirement else "None"
                linked_tcs = ", ".join([tc.test_id for tc in task.test_cases.filter(is_in_trash=False)]) or "None"
                due_date_str = task.due_date.strftime('%Y-%m-%d') if task.due_date else "N/A"
                lines.append(f"| {task.task_id} | {task.title} | {linked_req} | {assignees} | {task.get_status_display()} | {task.get_priority_display()} | {linked_tcs} | {due_date_str} |")
            lines.append("")
            
        # 5. QA & Test Cases
        if test_cases.exists():
            lines.append("## 5. QA & Test Validation")
            lines.append("| TEST ID | Scenario/Title | Linked Task | Linked REQ | Priority | Status | Expected Result |")
            lines.append("| :---: | :--- | :---: | :---: | :---: | :---: | :--- |")
            for tc in test_cases:
                linked_task = tc.task.task_id if tc.task else "None"
                linked_req = tc.task.requirement.req_id if (tc.task and tc.task.requirement) else "None"
                expected_res = tc.expected_result or "None"
                lines.append(f"| {tc.test_id} | {tc.title} | {linked_task} | {linked_req} | {tc.priority.upper() if tc.priority else 'N/A'} | {tc.get_status_display()} | {expected_res} |")
            lines.append("")
            
        # 6. RTM
        if requirements.exists():
            lines.append("## 6. Requirement Traceability Matrix (RTM)")
            lines.append("| Requirement | Development Tasks | Test Cases |")
            lines.append("| :--- | :--- | :--- |")
            rtm_data = ReportEngine._get_rtm_data(requirements)
            for row in rtm_data:
                req_str = f"**{row['req'].req_id}**<br>{row['req'].name}"
                
                if row['tasks']:
                    task_items = []
                    for t in row['tasks']:
                        task_items.append(f"{t.task_id} ({t.get_status_display()})")
                    tasks_str = "<br>".join(task_items)
                else:
                    tasks_str = "*No tasks mapped*"
                    
                if row['test_cases']:
                    tc_items = []
                    for tc in row['test_cases']:
                        tc_items.append(f"{tc.test_id} ({tc.get_status_display()})")
                    tcs_str = "<br>".join(tc_items)
                else:
                    tcs_str = "*No verification coverage*"
                    
                lines.append(f"| {req_str} | {tasks_str} | {tcs_str} |")
            lines.append("")
            
        return "\n".join(lines).encode('utf-8')

    # ==========================================
    # FALLBACK / SPECIFIC REPORTS
    # ==========================================
    @staticmethod
    def generate_requirement_report(project, requirements, format='pdf', template_type='srs'):
        if format in ['pdf', 'html']:
            return ReportEngine._generate_master_pdf_html(project, requirements, [], [], format)
        elif format == 'docx':
            return ReportEngine._generate_master_docx(project, requirements, [], [])
        elif format == 'xlsx':
            return ReportEngine._generate_master_xlsx(project, requirements, [], [])
        elif format == 'md':
            return ReportEngine._generate_master_markdown(project, requirements, [], [])

    @staticmethod
    def generate_task_report(project, tasks, format='pdf'):
        if format in ['pdf', 'html']:
            return ReportEngine._generate_master_pdf_html(project, [], tasks, [], format)
        elif format == 'docx':
            return ReportEngine._generate_master_docx(project, [], tasks, [])
        elif format == 'xlsx':
            return ReportEngine._generate_master_xlsx(project, [], tasks, [])
        elif format == 'md':
            return ReportEngine._generate_master_markdown(project, [], tasks, [])

    @staticmethod
    def generate_test_report(project, test_cases, format='pdf'):
        if format in ['pdf', 'html']:
            return ReportEngine._generate_master_pdf_html(project, [], [], test_cases, format)
        elif format == 'docx':
            return ReportEngine._generate_master_docx(project, [], [], test_cases)
        elif format == 'xlsx':
            return ReportEngine._generate_master_xlsx(project, [], [], test_cases)
        elif format == 'md':
            return ReportEngine._generate_master_markdown(project, [], [], test_cases)
